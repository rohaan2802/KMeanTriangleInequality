#!/usr/bin/env python3
"""
Speaker notes for each slide of PDC_Kmeans_Presentation.pptx (order matches build_presentation.py).
Requires: python-docx
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_slide_section(doc: Document, n: int, title_on_slide: str, script: str) -> None:
    h = doc.add_heading(f"Slide {n} — {title_on_slide}", level=1)
    h.runs[0].font.size = Pt(14)
    for para in script.strip().split("\n\n"):
        p = doc.add_paragraph(para.strip())
        for run in p.runs:
            run.font.size = Pt(11)
    doc.add_paragraph()


def build_doc(out: Path) -> None:
    doc = Document()
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Presentation speaker notes\nPDC_Kmeans_Presentation.pptx")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph(
        "Slide numbering matches experiments/build_presentation.py (same order as the generated .pptx). "
        "Practice timing: aim ~45–60 seconds per content slide; figures ~60–90 seconds."
    )

    add_slide_section(
        doc,
        1,
        "Title",
        """Good morning/afternoon. We are presenting our Parallel and Distributed Computing final project on "
        "triangle-inequality K-means using the reference implementation from Kwedlo and Czochanski, IEEE Access 2019.\n\n"
        "We focus on two gaps on a single workstation under WSL: Gap 1 compares two OpenMP scheduling builds for Elkan—"
        "what we call static-oriented versus guided—while sweeping thread count. Gap 3 varies the feature dimension M "
        "on synthetic data and compares Lloyd versus Elkan using the program’s reported metrics.\n\n"
        "We did not rewrite the authors’ core clustering algorithms; our contribution is experimental control and analysis."""
    )

    add_slide_section(
        doc,
        2,
        "Deliverables & artifacts",
        """Brief orientation for graders: our written report is generated as PDC_Kmeans_Final_Report.docx from a script, "
        "then exported to PDF. This deck is built by build_presentation.py from plots under experiments/figures.\n\n"
        "Everything ties to a pinned upstream Git commit so reviewers can rebuild binaries and logs using README instructions."""
    )

    add_slide_section(
        doc,
        3,
        "Problem & motivation",
        """K-means repeatedly assigns points to nearest centroids and recomputes means. The naive Lloyd assignment "
        "does a lot of distance work—roughly O(N×K×M) per iteration—so it dominates runtime at scale.\n\n"
        "Elkan’s method uses bounds and the triangle inequality to skip distance evaluations when safe. "
        "The base paper parallelizes these algorithms with MPI and OpenMP and shows scheduling matters on clusters.\n\n"
        "We ask narrower workstation-scale questions about scheduling and about how dimensionality affects pruning."""
    )

    add_slide_section(
        doc,
        4,
        "Base paper (short)",
        """The paper implements multiple algorithms—Lloyd, Elkan, Annulus, Drake, Yinyang—with hybrid MPI plus OpenMP, "
        "uses data parallelism in assignment, and compares static versus guided OpenMP on large hardware.\n\n"
        "We deliberately use the single-node OpenMP kmeans binary from the same repository—not the full MPI driver—"
        "so our experiments fit one laptop and stay reproducible for the course timeline."""
    )

    add_slide_section(
        doc,
        5,
        "Project scope — two gaps",
        """Gap 1: we built two executables from the same sources—default OpenMP loops versus Makefile DYNAMIC=y for guided-style "
        "scheduling—and ran Elkan on one medium dataset while varying OMP_NUM_THREADS.\n\n"
        "Gap 3: we generated synthetic binaries for several M values at fixed N and K, eight threads, and recorded "
        "wall time and the average distance-calculation ratio for naive versus Elkan.\n\n"
        "Out of scope: multi-node MPI, GPU ports, and editing the authors’ core C++ algorithms."""
    )

    add_slide_section(
        doc,
        6,
        "Methodology",
        """We pinned the Bitbucket hybrid-triangle-kmeans commit shown on the slide. No patches to Clust/*.cpp.\n\n"
        "The only difference between our two binaries is Makefile flags: OPT=y OPENMP=y versus adding DYNAMIC=y. "
        "We copy the produced binaries to experiments/bin as kmeans_static_gap1 and kmeans_guided_gap1.\n\n"
        "We parse stdout with parse_gap1_log.py and label pairs OK when iterations and SSE match; otherwise DIFF_PATH—"
        "meaning a scheduling-only comparison is not fair."""
    )

    add_slide_section(
        doc,
        7,
        "Experimental setup",
        """Environment: WSL Ubuntu, g++, OpenMP via -fopenmp. We fixed Forgy initialization, random seed 42, "
        "relative convergence tolerance 1e-4, and K equals 256 for our main tables.\n\n"
        "Gap 1 logs go to gap1_medium_runs.txt; Gap 3 uses run_gap3_matrix.sh and gap3_runs.txt. "
        "Figures come from plot_pdc_figures.py."""
    )

    add_slide_section(
        doc,
        8,
        "Live demo",
        """If asked to demonstrate live, we export OMP_NUM_THREADS, point to a small or medium .bin under data_gap3, "
        "and run kmeans_static_gap1 with -a elkan -c 256 -r 42 -R 1e-4.\n\n"
        "We highlight two stdout lines: k-means execution time and Avg distance calculations ratio. "
        "Optionally repeat with kmeans_guided_gap1 for contrast—while noting comparability caveats."""
    )

    add_slide_section(
        doc,
        9,
        "Gap 1 figure — Comparable pairs",
        """This plot shows median wall-clock time for pairs where static and guided runs matched iterations and SSE—"
        "so we can interpret scheduling fairly. You typically see this for lower thread counts in our study.\n\n"
        "Walk the audience along the x-axis: threads 1 and 2. State whether guided is slightly faster or similar "
        "when comparable—quote approximate medians from your table if asked."""
    )

    add_slide_section(
        doc,
        10,
        "Gap 1 figure — All medians",
        """This bar chart includes all paired runs even when paths differ—so T equals 4 and 8 still appear.\n\n"
        "Emphasize: raw bars can mislead when DIFF_PATH is common—those points are not apples-to-apples scheduling wins.\n\n"
        "Use this slide to motivate why we require OK pairs for strict scheduling conclusions."""
    )

    add_slide_section(
        doc,
        11,
        "Gap 1 — Interpretation",
        """Summarize: at one and two threads many pairs were comparable; guided could be modestly faster.\n\n"
        "At four and eight threads we often saw DIFF_PATH—different iteration counts or SSE between builds—because "
        "parallel floating-point reduction order and schedule interact with Elkan’s branching structure.\n\n"
        "Fair takeaway: scheduling effects exist but must be reported with comparability checks—not a single bare winner."""
    )

    add_slide_section(
        doc,
        12,
        "Gap 3 — Ratio vs M",
        """This curve plots the program’s average distance-calculation ratio versus dimension M for Elkan.\n\n"
        "Explain intuition: in higher dimensions, distances concentrate; bounds become less discriminating, "
        "so pruning weakens and the ratio trends upward toward full-work behavior.\n\n"
        "Mention our data is uniform random in [0,1)—a limitation, but the trend matches theory discussion."""
    )

    add_slide_section(
        doc,
        13,
        "Gap 3 — Time vs M",
        """Overlay naive Lloyd versus Elkan static and guided runtimes across M.\n\n"
        "At low M, Elkan is dramatically faster than naive on our generator; the gap narrows as M grows.\n\n"
        "If guided versus static differs at some M, tie back to SSE comparability columns in the report when needed."""
    )

    add_slide_section(
        doc,
        14,
        "Gap 3 — Speedup naive/Elkan",
        """This line shows naive wall time divided by Elkan static time—an algorithmic speedup proxy at fixed thread count.\n\n"
        "It reinforces that Elkan’s advantage is largest at small M for our synthetic setup.\n\n"
        "Caution: speedup is not parallel speedup versus one thread unless you measured that separately."""
    )

    add_slide_section(
        doc,
        15,
        "Gap 3 — Interpretation",
        """The ratio rises with M; Elkan’s runtime advantage over naive shrinks as pruning weakens.\n\n"
        "Stress limitations: one run per configuration in our matrix, i.i.d. uniform—not clustered real-world features.\n\n"
        "Connect briefly to distance concentration literature if panel asks ‘why’."""
    )

    add_slide_section(
        doc,
        16,
        "Limitations & reproducibility",
        """WSL introduces timing jitter; ignore noisy thread-clock diagnostics.\n\n"
        "Single runs per Gap 3 cell—medians across repetitions would be stronger science.\n\n"
        "Anyone can regenerate logs and figures from README with the same commit hash and scripts—we value reproducibility."""
    )

    add_slide_section(
        doc,
        17,
        "Conclusion",
        """Closing message: scheduling impact on Elkan is thread-dependent and often requires comparability filtering.\n\n"
        "Dimensionality strongly affects pruning and therefore Elkan’s edge over Lloyd on our synthetic sweep.\n\n"
        "Future work: repeated trials, clustered mixtures, deterministic reductions, optional MPI scaling."""
    )

    add_slide_section(
        doc,
        18,
        "LLM / AI assistance",
        """Course disclosure: we used LLM assistants for outlining, Makefile explanations, and scripting support.\n\n"
        "All binaries were executed locally; we verified iterations and SSE where claiming OK pairs.\n\n"
        "Both group members understand stdout fields, plots, and limitations."""
    )

    add_slide_section(
        doc,
        19,
        "Thank you — Questions?",
        """Thank the panel. Offer to show raw logs or rerun a command if helpful.\n\n"
        "Anticipate questions on DIFF_PATH, why not MPI, .bin format, and choice of K and stopping tolerance.\n\n"
        "End confidently and invite technical questions."""
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved: {out}")


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--out",
        type=Path,
        default=Path.home() / "22i-2327_F_FinalProject" / "PDC_Kmeans_Slide_Speaker_Notes.docx",
    )
    args = ap.parse_args()
    build_doc(args.out.expanduser().resolve())


if __name__ == "__main__":
    main()
