#!/usr/bin/env python3
"""
Generate demo/defense Q&A as Word .docx (python-docx).
Run on WSL: source .venv/bin/activate && pip install python-docx && python3 experiments/build_demo_prep_docx.py
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def qa(doc: Document, question: str, answer: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Q: ")
    r.bold = True
    p.add_run(question)
    a = doc.add_paragraph()
    r2 = a.add_run("A: ")
    r2.bold = True
    a.add_run(answer)
    doc.add_paragraph()


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--out",
        type=Path,
        default=Path.home() / "22i-2327_F_FinalProject" / "PDC_Kmeans_Demo_Defense_QA.docx",
    )
    ap.add_argument(
        "--project-root",
        type=Path,
        default=Path.home() / "22i-2327_F_FinalProject",
    )
    args = ap.parse_args()
    root = args.project_root.expanduser().resolve()
    out = args.out.expanduser().resolve()

    doc = Document()
    doc.sections[0].top_margin = Inches(1)
    doc.sections[0].bottom_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(
        "PDC Final Project — Demo & Defense Preparation\n"
        "K-means / Triangle Inequality / OpenMP Scheduling vs Dimensionality\n\n"
        "Question–Answer Reference (study guide)"
    )
    tr.bold = True
    tr.font.size = Pt(14)
    doc.add_paragraph(
        "Use this document to prepare for live demo and oral defense. "
        "Adapt answers to your exact commands and machine; commit hash and paths assume "
        "~/22i-2327_F_FinalProject on WSL."
    )

    # --- Cheat sheet ---
    doc.add_heading("0. One-minute cheat sheet", level=1)
    cheat = (
        "• Base paper: Kwedlo & Czochanski (IEEE Access 2019), hybrid MPI/OpenMP K-means with triangle inequality.\n"
        "• Our scope: author code unchanged; single-node OpenMP on WSL; two experiments.\n"
        "• Gap 1: static-oriented build vs guided build (Makefile DYNAMIC=y) for Elkan; vary OMP_NUM_THREADS; "
        "dataset gap1_medium (N=200k, M=128, K=256).\n"
        "• Gap 3: synthetic bins N=100k, K=256, T=8; M ∈ {2,8,16,32,64,128}; naive vs Elkan; ratio vs time.\n"
        "• Fairness: compare scheduling only when iterations + SSE match (OK); else DIFF_PATH.\n"
        "• Tools: g++/OpenMP, make, Python (parse/plot/build_report/build_slides), matplotlib, python-pptx, python-docx.\n"
        "• Upstream commit: 0265d45a2eb04fec01ed53fc8635b277082ee284."
    )
    doc.add_paragraph(cheat)

    doc.add_heading("1. What is this project about?", level=1)
    qa(
        doc,
        "What problem did you study?",
        "Parallel k-means clustering. Assignment (nearest centroid) dominates cost; triangle-inequality variants "
        "like Elkan prune distance computations. The base paper scales hybrid MPI/OpenMP on clusters; we narrow "
        "to one workstation and study (1) OpenMP scheduling for Elkan and (2) how pruning/runtime depend on "
        "feature dimension M on synthetic data.",
    )
    qa(
        doc,
        "Why is k-means expensive?",
        "Each iteration does assignment then centroid update. Lloyd’s naive assignment is O(N·K·M) distance work "
        "per iteration in the worst case when every point compares to every centroid every time.",
    )
    qa(
        doc,
        "What is triangle inequality in this context?",
        "Using bounds (upper/lower) on distances so we can prove some centroid pairs cannot be nearest and skip "
        "explicit distance computations. Elkan’s algorithm maintains such bounds; it stays exact relative to "
        "the same numerical path as the implementation’s Lloyd baseline within floating-point behavior.",
    )

    doc.add_heading("2. Base paper & your relation to it", level=1)
    qa(
        doc,
        "Summarize the base paper in one paragraph.",
        "Kwedlo & Czochanski implement Lloyd and several accelerated variants (Elkan, Annulus, Drake, Yinyang) "
        "with hybrid MPI/OpenMP. Work is partitioned across MPI ranks; updates use reductions. They compare "
        "OpenMP scheduling policies on large hardware and use a high-dimensional descriptor dataset.",
    )
    qa(
        doc,
        "Did you use MPI?",
        "No for our timed experiments. We used the authors’ single-node OpenMP kmeans binary. MPI scaling was "
        "out of scope.",
    )
    qa(
        doc,
        "Did you modify their C++ clustering code?",
        "No. We built two binaries from the same sources using different Makefile flags (OPENMP=y vs DYNAMIC=y). "
        "Our contribution is experimental design, logs, parsing, plots, report, and slides—not a new algorithm.",
    )

    doc.add_heading("3. Gap 1 — OpenMP scheduling (static-oriented vs guided)", level=1)
    qa(
        doc,
        "What exactly did you compare?",
        "Two builds of the same program: (1) default OpenMP loop scheduling as built by "
        "`make kmeans OPT=y OPENMP=y`, copied to kmeans_static_gap1; (2) guided scheduling enabled via "
        "`make … DYNAMIC=y`, copied to kmeans_guided_gap1. Both run Elkan with identical CLI flags on the "
        "same .bin file.",
    )
    qa(
        doc,
        "What does DYNAMIC=y do in this codebase?",
        "It defines OMPDYNAMIC as schedule(guided) on annotated `#pragma omp for` loops (see upstream Makefile / "
        "headers). So it is not toggling OMP_DYNAMIC at runtime—it changes compile-time scheduling on specific loops.",
    )
    qa(
        doc,
        "Why might guided scheduling help Elkan?",
        "Elkan’s per-point work is uneven: some points skip many distances; others do not. Static partitioning "
        "of loop iterations can leave threads idle (tail imbalance). Guided scheduling hands out shrinking chunks "
        "so threads can steal smaller pieces of remaining work—often helping dynamic imbalance.",
    )
    qa(
        doc,
        "Why might guided hurt?",
        "More scheduling overhead; worse cache/locality vs static chunks; NUMA effects (paper discusses "
        "memory bandwidth). On small thread counts the tradeoff differs from a 64-node cluster.",
    )
    qa(
        doc,
        "What dataset and parameters did you use for Gap 1?",
        "gap1_medium.bin with N=200000, M=128; K=256; Forgy init; seed -r 42; stopping -R 1e-4; algorithm "
        "-a elkan. Threads: typically 1,2,4,8 via OMP_NUM_THREADS.",
    )
    qa(
        doc,
        "What is OK vs DIFF_PATH?",
        "We pair consecutive STATIC and GUIDED runs with the same thread count and repeat index. OK means same "
        "iteration count and SSE bracket total within tolerance (~0.5). DIFF_PATH means the numerical trajectory "
        "differed—often due to parallel reduction order plus scheduling—so a pure scheduling speed claim is weak.",
    )
    qa(
        doc,
        "Your Gap 1 table shows no comparable pairs at T=4,8. What do you conclude?",
        "We report medians but emphasize we cannot fairly compare scheduling-only speedups when paths diverge. "
        "This is an honest experimental outcome, not a bug.",
    )

    doc.add_heading("4. Gap 3 — Dimensionality (M sweep)", level=1)
    qa(
        doc,
        "What was the research question?",
        "How Elkan’s pruning (Avg distance calculations ratio) and runtime versus naive Lloyd change as feature "
        "dimension M increases on synthetic uniform data, holding N, K, threads fixed.",
    )
    qa(
        doc,
        "How did you generate datasets?",
        "Python script gen_gap3_bins.py writes author-format binaries: int32 N, int32 M, float32 row-major data. "
        "We used fixed seed (e.g. 42) for reproducibility.",
    )
    qa(
        doc,
        "What does ‘Avg distance calculations ratio’ mean?",
        "It is reported by the reference implementation as a percentage-like workload measure relative to full "
        "work; naive shows 100; Elkan shows lower values when pruning works (fewer effective distance evaluations).",
    )
    qa(
        doc,
        "Why does ratio often increase with M?",
        "Informally: in higher dimensions distances concentrate (nearest vs farthest ratios shrink), bounds "
        "become tight less often, pruning weakens—consistent with classical ‘curse of dimensionality’ discussions "
        "(cite Beyer/Aggarwal in report). Synthetic uniform data is not clustered; real clustering workloads may differ.",
    )
    qa(
        doc,
        "Why compare naive vs Elkan times?",
        "To show algorithmic advantage of triangle inequality when pruning is strong at low M and how it erodes as M grows.",
    )

    doc.add_heading("5. Tools, environment, and reproducibility", level=1)
    qa(
        doc,
        "What OS and hardware did you use?",
        "Document what you actually used—example from your logs: WSL2 Linux on x86_64; Intel i5-5300U; ~8 GB RAM; "
        "8 logical CPUs. Update if you demo on another machine.",
    )
    qa(
        doc,
        "Which compiler and flags?",
        "g++ per upstream Options-x86_64.gcc / Makefile: optimization enabled for release kmeans, -fopenmp for OpenMP.",
    )
    qa(
        doc,
        "Which Python packages?",
        "matplotlib for plots; python-pptx for slides; python-docx for Word outputs; standard library for parsers. "
        "Install in a venv on Ubuntu (PEP 668 blocks global pip without --break-system-packages).",
    )
    qa(
        doc,
        "How do you reproduce figures?",
        "Ensure gap1_medium_runs.txt and gap3_runs.txt exist under experiments/, then run plot_pdc_figures.py "
        "which writes PNGs and CSV/TABLE markdown under experiments/figures/.",
    )
    qa(
        doc,
        "How do you regenerate report and slides?",
        "build_final_report_docx.py → PDC_Kmeans_Final_Report.docx; build_presentation.py → PDC_Kmeans_Presentation.pptx.",
    )

    doc.add_heading("6. Live demo — what can you show?", level=1)
    qa(
        doc,
        "Give a safe 2–3 minute demo script.",
        "1) Show dataset path: ls -lh experiments/data_gap3/*.bin. 2) export OMP_NUM_THREADS=4. "
        "3) Run kmeans_static_gap1 on a small/medium .bin with -a elkan -c 256 -v 1 -r 42 -R 1e-4. "
        "4) Point at stdout: k-means execution time, total iterations, Best MSE (SSE), Avg distance calculations ratio. "
        "5) Optionally repeat with kmeans_guided_gap1 and mention comparability caveats.",
    )
    qa(
        doc,
        "What if the demo run is slow?",
        "Pick the smallest binary you have (e.g. tiny.bin or smallest gap3 M) or reduce verbosity; avoid huge K/N "
        "if time-limited.",
    )

    doc.add_heading("7. Limitations & honesty", level=1)
    qa(
        doc,
        "What are the main limitations?",
        "WSL timing noise; synthetic i.i.d. data may not reflect real cluster geometry; single-run configs for Gap 3 "
        "unless you repeated medians; no MPI scaling; no GPU; schedule comparisons confounded when DIFF_PATH occurs.",
    )
    qa(
        doc,
        "Did you prove guided is always faster?",
        "No. We report medians and explicitly restrict strong claims to comparable pairs.",
    )

    doc.add_heading("8. Academic integrity & LLM use", level=1)
    qa(
        doc,
        "How did you use AI assistants?",
        "Per course disclosure: planning, scripting help, README/report structure, environment troubleshooting. "
        "All binaries were built locally; logs and numbers come from your runs; both members must explain results.",
    )

    doc.add_heading("9. Hard questions — short answers", level=1)
    qa(
        doc,
        "Is Elkan always faster than Lloyd?",
        "Not guaranteed—especially when pruning is weak (high M) or bound maintenance overhead dominates small cases.",
    )
    qa(
        doc,
        "Why not SIMD?",
        "Explicit intrinsics were out of one-week scope; compiler auto-vectorization may still apply via flags.",
    )
    qa(
        doc,
        "What Bottleneck: compute vs memory?",
        "Naive is often compute-heavy distance work; Elkan shifts work but adds memory for bounds. "
        "High-K Elkan can be memory-heavy (paper notes). Tie answers to your stdout + optional perf stat.",
    )
    qa(
        doc,
        "What is parallel efficiency?",
        "Speedup vs threads divided by thread count; we focused more on scheduling comparability and pruning metrics "
        "than cluster-wide efficiency—state if asked.",
    )

    doc.add_heading("10. Dataset files & binary layout", level=1)
    qa(
        doc,
        "What format are the .bin datasets?",
        "Per upstream reader: int32 N, int32 M, then N×M float32 values row-major (each row is one point). "
        "Little-endian on typical x86.",
    )
    qa(
        doc,
        "Where do gap3 files live?",
        "experiments/data_gap3/gap3_N100000_M*.bin generated by gen_gap3_bins.py.",
    )
    qa(
        doc,
        "Why synthetic data instead of GIST/Tiny Images?",
        "Controlled sweep over M with fixed N,K,seed; matches project scope. Paper notes low-D experiments could differ—we cite that as motivation.",
    )

    doc.add_heading("11. Scripts you used — what do they do?", level=1)
    qa(
        doc,
        "What does parse_gap1_log.py output?",
        "Reads gap1_medium_runs.txt, finds paired STATIC/GUIDED blocks, prints per-thread comparable counts, "
        "median times, median guided/static ratio when comparable, and labels each pair OK or DIFF_PATH.",
    )
    qa(
        doc,
        "What does plot_pdc_figures.py produce?",
        "Parses Gap 1 and Gap 3 logs, writes summary CSVs, PNG figures (Gap 1 medians, Gap 3 ratio/time/speedup), "
        "and TABLE_A/TABLE_B markdown for the report.",
    )
    qa(
        doc,
        "What does run_gap3_matrix.sh do?",
        "Runs naive + Elkan static + Elkan guided across generated binaries (your version may loop M values); "
        "stdout appended to gap3_runs.txt.",
    )

    doc.add_heading("12. Understanding stdout (noise vs signal)", level=1)
    qa(
        doc,
        "What is ‘Highest absolute difference between thread 0 and real time’?",
        "An internal timing diagnostic from the program; can look alarming under WSL. Use wall k-means execution time "
        "as the primary timing metric unless you deeply validated thread timers.",
    )
    qa(
        doc,
        "What does ‘NUMA optimizations on’ mean?",
        "Upstream enables NUMA-aware allocation paths when supported. On a typical laptop/WSL single-socket setup "
        "this may have limited effect—mention without overstating.",
    )

    doc.add_heading("13. Metrics & definitions", level=1)
    qa(
        doc,
        "How do you define speedup for Gap 3?",
        "Common presentation: naive_time / elkan_time for the same dataset and CLI—algorithmic speedup of Elkan "
        "over Lloyd for that run. Parallel speedup would vary OMP threads separately.",
    )
    qa(
        doc,
        "Why fix OMP_NUM_THREADS=8 for Gap 3?",
        "Isolates dimensionality effect by holding parallel configuration constant.",
    )
    qa(
        doc,
        "What is SSE / MSE here?",
        "The program prints Best MSE (SSE) with a normalized and bracket total form—use bracket totals for "
        "comparability checks between runs.",
    )

    doc.add_heading("14. Comparison to the base paper’s findings", level=1)
    qa(
        doc,
        "Does your Gap 1 contradict the paper?",
        "Not necessarily—the paper’s scheduling results are on MPI+many cores and Elkan memory behavior differs. "
        "Our workstation/WSL study is a different regime; partial comparability (DIFF_PATH) is expected.",
    )

    doc.add_heading("15. What you did NOT do (say clearly if asked)", level=1)
    doc.add_paragraph(
        "• Multi-node MPI scaling experiments.\n"
        "• GPU / CUDA / OpenCL implementation.\n"
        "• Changing Elkan’s numerical core or adding new pruning rules.\n"
        "• Full thread×K×N factorial matrix at cluster scale.\n"
        "• Running on the paper’s exact public dataset (optional extension)."
    )

    doc.add_heading("16. Instructor-style rapid-fire", level=1)
    qa(doc, "Name the base paper and venue.", "Kwedlo & Czochanski, IEEE Access, 2019 (DOI in README).")
    qa(doc, "What is Elkan’s algorithm known for?", "Triangle inequality bounds to skip distance computations in exact k-means assignment steps (within FP behavior).")
    qa(doc, "What is Lloyd’s algorithm?", "Classic alternate assignment to nearest centroid and recomputation of means until convergence criterion.")
    qa(doc, "Why OpenMP?", "Shared-memory parallelism over the assignment/update loops in the single-node build.")
    qa(doc, "What is schedule(guided)?", "OpenMP guided scheduling: chunk sizes decrease over time to reduce tail imbalance in parallel for loops.")
    qa(doc, "Why might iterations differ between binaries?", "Floating-point non-associativity + parallel reductions + different iteration ordering can change when stopping criterion triggers.")

    doc.add_heading("17. File & command reference", level=1)
    doc.add_paragraph(
        f"Project root (yours): {root}\n"
        "Key binaries: experiments/bin/kmeans_static_gap1, kmeans_guided_gap1\n"
        "Key logs: experiments/gap1_medium_runs.txt, experiments/gap3_runs.txt\n"
        "Figures: experiments/figures/*.png\n"
        "Commit: 0265d45a2eb04fec01ed53fc8635b277082ee284"
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print("Saved:", out)


if __name__ == "__main__":
    main()
