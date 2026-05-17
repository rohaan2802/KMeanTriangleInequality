#!/usr/bin/env python3
"""
Build demo / defense Q&A as Word .docx. Requires: python-docx (same venv as other tools).
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
def add_qa(doc: Document, question: str, answer: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Q: ")
    r.bold = True
    p.add_run(question)
    a = doc.add_paragraph()
    a.add_run("A: ").bold = True
    a.add_run(answer)
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def build_doc(out: Path) -> None:
    doc = Document()
    doc.sections[0].page_width = Inches(8.5)
    doc.sections[0].page_height = Inches(11)
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "PDC Final Project — Demo & Defense Q&A\n"
        "K-means (triangle inequality) — scheduling & dimensionality study"
    )
    r.bold = True
    r.font.size = Pt(15)
    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(
        "Use with: hybrid-triangle-kmeans (Kwedlo & Czochanski, IEEE Access 2019), "
        "commit 0265d45a2eb04fec01ed53fc8635b277082ee284. "
        "Workstation / WSL2; no edits to author Clust/*.cpp."
    )

    add_h(doc, "1. One-minute project summary (elevator pitch)", 1)
    add_qa(
        doc,
        "What is this project about in one sentence?",
        "We take the authors’ public OpenMP k-means code and run a controlled performance study on one machine: "
        "Gap 1 compares two OpenMP scheduling builds (static-oriented vs guided) for Elkan; "
        "Gap 3 varies data dimension M on synthetic data and measures Lloyd vs Elkan and the program’s "
        "“average distance calculations” ratio. We do not change the core algorithms.",
    )
    add_qa(
        doc,
        "What problem does K-means solve?",
        "Partition N points in M dimensions into K clusters by alternating (1) assign each point to the nearest "
        "centroid and (2) recompute centroids as means of assigned points, until convergence or max iterations.",
    )
    add_qa(
        doc,
        "Why is Lloyd’s method expensive?",
        "In the naive form, each point can need distance to all K centroids every iteration — O(N·K·M) work per "
        "iteration in the worst case. That dominates when N and K are large.",
    )
    add_qa(
        doc,
        "What is Elkan’s contribution here?",
        "Elkan’s algorithm uses the triangle inequality and bounds (upper/lower) to skip many distance "
        "evaluations while still finding the same assignment as full distance search under the same floating-point "
        "model in the given implementation. The program reports an “Avg distance calculations ratio” — lower "
        "usually means more pruning.",
    )

    add_h(doc, "2. Base paper (Kwedlo & Czochanski, 2019)", 1)
    add_qa(
        doc,
        "What does the base paper do that we did not fully reproduce?",
        "The paper implements hybrid MPI + OpenMP for several triangle-inequality algorithms and runs at cluster "
        "scale with a high-dimensional image descriptor dataset, comparing static vs guided OpenMP scheduling. "
        "We use the same single-node kmeans binary, not the multi-node MPI driver, and we use synthetic + our own "
        "medium dataset instead of their full distributed setup.",
    )
    add_qa(
        doc,
        "Why is guided scheduling discussed in the paper?",
        "Triangle-inequality assignment can create uneven work per point (some points prune heavily, others "
        "almost not). Static chunking can leave threads idle at barriers (load imbalance). Guided scheduling "
        "can reduce tail imbalance by handing smaller chunks later in the loop — but it adds overhead and can "
        "hurt locality/NUMA patterns depending on the algorithm and hardware.",
    )
    add_qa(
        doc,
        "What did we borrow vs what is our contribution?",
        "Borrowed: all clustering implementation from upstream. Ours: experimental matrix (thread sweep + "
        "dimension sweep), parsing comparable vs non-comparable runs (DIFF_PATH), plots/tables, interpretation.",
    )

    add_h(doc, "3. Gap 1 — Static-oriented vs guided scheduling", 1)
    add_qa(
        doc,
        "What exactly did you build?",
        "Two executables from the same sources: (1) make kmeans OPT=y OPENMP=y → copied to "
        "experiments/bin/kmeans_static_gap1; (2) make … OPENMP=y DYNAMIC=y → kmeans_guided_gap1. "
        "DYNAMIC=y injects schedule(guided) on annotated loops via Makefile macros.",
    )
    add_qa(
        doc,
        "What input did you use for Gap 1?",
        "A medium binary dataset gap1_medium.bin with N=200000, M=128; runs use K=256, Forgy init, seed -r 42, "
        "relative stop -R 1e-4, algorithm -a elkan.",
    )
    add_qa(
        doc,
        "Which threads did you test?",
        "OMP_NUM_THREADS ∈ {1, 2, 4, 8} with multiple paired STATIC/GUIDED runs logged to gap1_medium_runs.txt.",
    )
    add_qa(
        doc,
        "What does OK vs DIFF_PATH mean in your parser?",
        "For each pair of STATIC and GUIDED runs with the same thread count and repeat index, we compare total "
        "iterations and the bracket total SSE printed by the program. If both match within tolerance (~0.5 on SSE), "
        "the pair is OK for a scheduling-only time comparison. If iterations or SSE differ, we label DIFF_PATH — "
        "the numerical trajectory changed (common with parallel reductions and schedule changes).",
    )
    add_qa(
        doc,
        "Why did DIFF_PATH appear at higher thread counts?",
        "Floating-point addition order differs across threads/chunks; tiny drift can change which bound tests "
        "fire early in Elkan, changing iteration counts or terminal SSE slightly. Then comparing raw wall times "
        "as “who won scheduling” is misleading.",
    )
    add_qa(
        doc,
        "What did you conclude for Gap 1?",
        "At T=1–2 many pairs were comparable; guided could be slightly faster when comparable. At T=4–8 we often "
        "saw no comparable pairs — report cautiously; emphasize methodology (fair comparison only when paths match).",
    )

    add_h(doc, "4. Gap 3 — Dimensionality (M) vs pruning and runtime", 1)
    add_qa(
        doc,
        "What datasets did you use?",
        "Synthetic binaries gap3_N100000_M{2,8,16,32,64,128}.bin from gen_gap3_bins.py: Python rng.random() floats "
        "in [0,1), fixed seed --seed 42. Header int32 N, M then row-major float32. Fixed N=100000, K=256, OMP_NUM_THREADS=8.",
    )
    add_qa(
        doc,
        "What algorithms did you run per file?",
        "Typically naive (-a naive) for Lloyd baseline, then Elkan with static binary and Elkan with guided binary — "
        "as captured in gap3_runs.txt / run_gap3_matrix.sh.",
    )
    add_qa(
        doc,
        "Why does the distance-calculation ratio increase with M?",
        "On uniform random data, distances concentrate in high dimensions (nearest vs farthest neighbors become "
        "similar). Triangle inequality filters need separation between centroid distances; when separation shrinks, "
        "pruning weakens — ratio trends toward 100 (full work).",
    )
    add_qa(
        doc,
        "Why might naive vs Elkan SSE rows say ‘differs’ in your table?",
        "Floating paths can diverge slightly; also Lloyd vs Elkan may exit with tiny SSE differences even when "
        "assignments are effectively aligned — always relate claims to your tolerance and logs.",
    )

    add_h(doc, "5. Tools, environment, and reproducibility", 1)
    add_qa(
        doc,
        "What OS and hardware did you use?",
        "Development/build on WSL2 Ubuntu; CPU/RAM as recorded in the report (example: Intel i5-5300U, ~8 GB RAM, "
        "8 logical CPUs). WSL adds timing jitter — mention when asked about noise.",
    )
    add_qa(
        doc,
        "Which compilers and flags matter?",
        "g++ with -fopenmp per upstream Makefile; release targets use aggressive optimization (-Ofast-style) and "
        "OpenMP as shipped by the authors — see Options-x86_64.gcc / Makefile.",
    )
    add_qa(
        doc,
        "Which Python tools did you use?",
        "matplotlib for plots (plot_pdc_figures.py); parse_gap1_log.py for Gap 1 summaries; python-pptx for slides "
        "(build_presentation.py); python-docx for report (build_final_report_docx.py). Use a venv if pip is PEP 668.",
    )
    add_qa(
        doc,
        "Did you use a profiler like VTune or perf?",
        "Primary metrics came from program stdout (wall time, ratio, iterations, SSE). Optional: perf stat for "
        "hardware counters — only cite if you actually ran it.",
    )
    add_qa(
        doc,
        "Where are logs and figures?",
        "experiments/gap1_medium_runs.txt, experiments/gap3_runs.txt; summaries under experiments/figures/ after "
        "running plot_pdc_figures.py.",
    )

    add_h(doc, "6. Correctness and integrity", 1)
    add_qa(
        doc,
        "Did you modify the authors’ algorithms?",
        "No changes to the core C++ clustering logic for our claimed baseline; we only built two scheduling variants "
        "and ran experiments.",
    )
    add_qa(
        doc,
        "How do you know outputs are plausible?",
        "SSE decreases over iterations in healthy runs; iteration counts are finite; naive ratio reads 100 in logs; "
        "Elkan ratio < 100 when pruning works; cross-check Elkan vs naive trends by dimension.",
    )
    add_qa(
        doc,
        "How did you use LLMs if asked?",
        "Support only — outlining, environment fixes, plotting workflow ideas. All binaries run locally; both "
        "members understand stdout fields and figures.",
    )

    add_h(doc, "7. Limitations (expect these questions)", 1)
    add_qa(
        doc,
        "What are the main limitations?",
        "Single workstation not cluster; WSL timing noise; synthetic i.i.d. data not real vision features; "
        "single run per Gap 3 cell in the matrix unless you repeated; DIFF_PATH complicates pure scheduling claims.",
    )
    add_qa(
        doc,
        "What would you do with more time?",
        "Multiple trials with medians; clustered synthetic mixtures; investigate deterministic reductions; optional "
        "MPI scaling in lab; larger M grid.",
    )

    add_h(doc, "8. OpenMP, Makefile, and binary details", 1)
    add_qa(
        doc,
        "What does DYNAMIC=y do in the Makefile?",
        "It adds a compile-time define so relevant OpenMP for loops use schedule(guided) (or the macro the authors "
        "map to “dynamic/guided” style chunking) instead of the default static-style chunking in the default build.",
    )
    add_qa(
        doc,
        "Is the “static” binary truly static schedule everywhere?",
        "In our naming, “static” means the default upstream build without DYNAMIC=y — the actual schedule is "
        "whatever the author code uses for each loop (often static for regular parallel for). The important point "
        "is we compare two built artifacts: default vs DYNAMIC=y.",
    )
    add_qa(
        doc,
        "What is OMP_NUM_THREADS?",
        "It sets the number of OpenMP worker threads for parallel regions. We export it before each run to control "
        "the thread sweep in Gap 1 and fix T=8 in Gap 3.",
    )
    add_qa(
        doc,
        "What parallel pattern does k-means use here?",
        "Mainly data-parallel loops over points in assignment and reduction-style updates for centroids; the exact "
        "OpenMP usage is in the author source (KMAlgorithm / reducers).",
    )
    add_qa(
        doc,
        "Why do you see odd “thread 0 vs real time” lines in output?",
        "The program can print OpenMP / thread timing diagnostics; on WSL they may look noisy. Rely on "
        "k-means execution time for the main wall-clock comparison unless you instrumented more carefully.",
    )

    add_h(doc, "9. Data files and format", 1)
    add_qa(
        doc,
        "What is the .bin format?",
        "As in our README: int32 N, int32 M, then N×M float32 values row-major — matching the author loader.",
    )
    add_qa(
        doc,
        "Could you use real images instead of synthetic data?",
        "Yes with enough engineering (feature extraction). Outside scope; synthetic lets us sweep M systematically.",
    )

    add_h(doc, "10. Metrics — speedup, efficiency, ratio", 1)
    add_qa(
        doc,
        "Did you measure classical parallel speedup Sp = T1/Tp?",
        "Gap 1 compares scheduling variants at fixed threads more than ideal speedup vs sequential; interpret carefully "
        "when DIFF_PATH breaks comparability. Gap 3 emphasizes naive vs Elkan at fixed T=8 across M.",
    )
    add_qa(
        doc,
        "What does ‘Avg distance calculations ratio’ mean practically?",
        "It is the implementation’s reported fraction of distance work vs a full N×K baseline accounting — lower "
        "means fewer distance evaluations per iteration on average.",
    )

    add_h(doc, "11. Scope boundaries — why not X?", 1)
    add_qa(
        doc,
        "Why no MPI / multi-node results?",
        "Out of scope for the one-week plan; the course still rewards single-node OpenMP study. The base paper’s "
        "MPI story is context, not our experiment matrix.",
    )
    add_qa(
        doc,
        "Why no GPU or SIMD intrinsics?",
        "Time and risk: GPU would need a different port; intrinsics need heavy validation. We focused on "
        "scheduling and dimensionality on the author CPU code.",
    )
    add_qa(
        doc,
        "Why not change the C++ to fix DIFF_PATH?",
        "That would be a new research/implementation project; we study the author build as given and report where "
        "comparisons are fair.",
    )

    add_h(doc, "12. Group work and academic integrity", 1)
    add_qa(
        doc,
        "How should each member answer “what did you do?”",
        "Both should know: how to run binaries, read logs, explain OK/DIFF_PATH, explain plots, and one limitation. "
        "Split who wrote which script is secondary to shared understanding of results.",
    )
    add_qa(
        doc,
        "Is using the authors’ code allowed?",
        "Yes as long as you cite the paper and repository and clearly separate your experimental work from their "
        "implementation.",
    )

    add_h(doc, "13. Live demo — what to show", 1)
    add_qa(
        doc,
        "What command would you run live?",
        "export OMP_NUM_THREADS=8; experiments/bin/kmeans_static_gap1 experiments/data_gap3/gap3_N100000_M32.bin "
        "-a elkan -c 256 -v 1 -r 42 -R 1e-4 — point to k-means execution time and Avg distance calculations ratio.",
    )
    add_qa(
        doc,
        "What if the demo machine differs from the report?",
        "Say clearly: numbers are from the logged environment; live run illustrates behavior, not identical seconds.",
    )

    add_h(doc, "14. Quick factual checklist", 1)
    p = doc.add_paragraph()
    p.add_run(
        "Upstream commit: 0265d45a2eb04fec01ed53fc8635b277082ee284 · "
        "Gap 1 binary names: kmeans_static_gap1, kmeans_guided_gap1 · "
        "Key logs: gap1_medium_runs.txt, gap3_runs.txt · "
        "K common: 256 · Stop: -R 1e-4 · Seed: -r 42 · "
        "Paper DOI: 10.1109/ACCESS.2019.2907885."
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved: {out}")


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--out",
        type=Path,
        default=Path.home() / "22i-2327_F_FinalProject" / "PDC_Kmeans_Demo_QA.docx",
    )
    args = ap.parse_args()
    build_doc(args.out.expanduser().resolve())


if __name__ == "__main__":
    main()
