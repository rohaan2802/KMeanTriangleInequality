#!/usr/bin/env python3
"""
Build PDC final report as Word .docx (tables, embedded plots, methodology).
Requires: pip install python-docx
"""
from __future__ import annotations

import argparse
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


def add_mono_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.rstrip() + "\n")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def parse_authors(project_root: Path) -> tuple[str, str]:
    p = project_root / "AUTHORS.md"
    if not p.is_file():
        return (
            "Mohammad Rohaan (22I-2327)",
            "[Partner full name] (22I-xxxx)",
        )
    t = p.read_text(encoding="utf-8", errors="replace")
    m1 = re.search(
        r"\|\s*Mohammad Rohaan\s*\|\s*(22I-\d+)\s*\|",
        t,
        re.IGNORECASE,
    )
    r1 = m1.group(1) if m1 else "22I-2327"
    return (
        f"Mohammad Rohaan ({r1})",
        "[Partner full name] (22I-xxxx) — replace before PDF",
    )


def try_shell_capture(cmd: list[str], timeout: int = 90) -> str | None:
    try:
        r = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if r.returncode != 0:
            return None
        return (r.stdout or "") + (r.stderr or "")
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return None


def build_doc(project_root: Path, out_path: Path) -> None:
    author1, author2 = parse_authors(project_root)
    exp = project_root / "experiments"
    fig_dir = exp / "figures"
    gap1_log = exp / "gap1_medium_runs.txt"
    gap3_log = exp / "gap3_runs.txt"

    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = Inches(1)
    sect.bottom_margin = Inches(1)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Parallel and Distributed Computing — Final Project\n"
        "K-means with Triangle Inequality: Scheduling vs Dimensionality"
    )
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(f"{author1}\n{author2}\nSection F")
    s.font.size = Pt(12)
    doc.add_paragraph()
    abs_p = doc.add_paragraph()
    abs_p.add_run("Abstract. ").bold = True
    abs_p.add_run(
        "We evaluate the public hybrid-triangle-kmeans implementation of Kwedlo & Czochanski "
        "(IEEE Access, 2019) on a single shared-memory machine (WSL2). "
        "Gap 1 compares default OpenMP loop scheduling versus a guided scheduling build (Makefile DYNAMIC=y) "
        "for Elkan’s algorithm while varying OMP_NUM_THREADS on a fixed medium dataset (N=200000, M=128, K=256). "
        "Gap 3 varies feature dimension M on synthetic datasets (N=100000, K=256, T=8), comparing Lloyd (naive) "
        "to Elkan and reporting the program’s average distance-calculation ratio and wall-clock time. "
        "We flag non-comparable scheduling pairs when iterations or SSE diverge (DIFF_PATH)."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "K-means alternates assignment (each point to nearest centroid) and update (centroid means). "
        "Lloyd’s assignment costs O(NKM) distance evaluations per iteration in the worst case. "
        "Triangle-inequality variants such as Elkan’s method prune distance computations while remaining exact "
        "under the same arithmetic model as implemented in the reference code."
    )
    doc.add_paragraph(
        "The base paper studies hybrid MPI/OpenMP implementations on a cluster and compares OpenMP scheduling "
        "policies at large scale. We narrow the scope to one workstation (WSL2) and two questions: "
        "(i) whether guided scheduling benefits Elkan at small thread counts, and "
        "(ii) how pruning effectiveness and runtime scale with dimensionality on synthetic data."
    )

    doc.add_heading("2. Base paper and problem context", level=1)
    doc.add_paragraph(
        "Kwedlo & Czochanski (IEEE Access, 2019) implement Lloyd and triangle-inequality variants (Elkan, Annulus, "
        "Drake, Yinyang) with hybrid MPI/OpenMP, partition data and bounds across ranks, and use global reductions "
        "in the update step. They compare static versus guided OpenMP scheduling on a large machine using a "
        "high-dimensional descriptor dataset."
    )
    doc.add_paragraph(
        "Our work reuses the authors’ single-node OpenMP kmeans binary at commit "
        "0265d45a2eb04fec01ed53fc8635b277082ee284 (Bitbucket: wkwedlo/hybrid-triangle-kmeans). "
        "We do not modify their clustering algorithms; contributions are controlled experiments and analysis."
    )

    doc.add_heading("3. Project scope and objectives", level=1)
    doc.add_heading("3.1 In scope", level=2)
    for x in (
        "Author kmeans: two builds — default OpenMP loops vs DYNAMIC=y (guided on annotated loops).",
        "Algorithms: -a naive (Lloyd) baseline; -a elkan for triangle-inequality study.",
        "Gap 1: OMP_NUM_THREADS ∈ {1,2,4,8}, dataset gap1_medium.bin (200000×128), K=256, seed 42, stop -R 1e-4.",
        "Gap 3: N=100000, K=256, T=8, M ∈ {2,8,16,32,64,128}, synthetic bins from gen_gap3_bins.py.",
    ):
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("3.2 Out of scope", level=2)
    for x in (
        "Multi-node MPI scaling (mpikmeans), GPU, edits to author Clust/*.cpp, hand-written SIMD intrinsics.",
    ):
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("3.3 Research questions", level=2)
    doc.add_paragraph(
        "RQ1: On WSL, does guided scheduling improve Elkan wall time versus the static-oriented build as threads "
        "increase, and when are runs not comparable because iterations or SSE differ?"
    )
    doc.add_paragraph(
        "RQ2: As M increases on i.i.d. synthetic data, how do the average distance-calculation ratio and "
        "Elkan vs naive runtime behave?"
    )

    doc.add_heading("4. Baseline method (Lloyd / naive)", level=1)
    doc.add_paragraph(
        "Baseline assignment evaluates distances to all K centroids per point (implementation: -a naive / NaiveKMA). "
        "Complexity of assignment is O(NKM) per iteration in the full-distance regime; the program reports "
        "Avg distance calculations ratio: 100 for naive runs in our logs."
    )

    doc.add_heading("5. Proposed methodology", level=1)
    doc.add_paragraph(
        "Elkan (-a elkan) maintains bounds to skip many distance evaluations. "
        "We compare two binaries: kmeans_static_gap1 (make OPT=y OPENMP=y) and kmeans_guided_gap1 "
        "(same with DYNAMIC=y). Fair comparison when total iterations and bracket SSE match; otherwise label "
        "DIFF_PATH (parallel floating-point reduction order and scheduling can change the numerical path)."
    )

    doc.add_heading("6. Experimental setup", level=1)
    doc.add_heading("6.1 Hardware and OS (recorded on WSL)", level=2)
    hw = doc.add_table(rows=1, cols=2)
    hw.style = "Table Grid"
    hdr = hw.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Value"
    for k, v in (
        ("OS / kernel", "Linux x86_64, WSL2 (example: 6.6.x-microsoft-standard-WSL2)"),
        ("CPU", "Intel Core i5-5300U @ 2.30GHz (replace if your machine differs)"),
        ("Logical CPUs", "8 (nproc)"),
        ("RAM (MemTotal)", "~7.9 GiB (8068228 kB — update if you rerun on another PC)"),
        ("Toolchain", "g++ (Ubuntu 13.3.x) — see README"),
    ):
        row = hw.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_heading("6.2 Software and baseline revision", level=2)
    doc.add_paragraph(
        "Compiler: g++ with -fopenmp per upstream Makefile; release kmeans targets. "
        "Git commit: 0265d45a2eb04fec01ed53fc8635b277082ee284."
    )

    doc.add_heading("6.3 Datasets and parameters", level=2)
    doc.add_paragraph(
        "Gap 1: gap1_medium.bin — N=200000, M=128. Gap 3: gap3_N100000_M*.bin under experiments/data_gap3/, "
        "seed 42 from gen_gap3_bins.py. Common: K=256, Forgy init (-t default), -r 42, -R 1e-4."
    )

    doc.add_heading("7. Measurement methodology (profiling)", level=1)
    doc.add_paragraph(
        "Primary metrics come from program stdout (application-level instrumentation), not an external sampling "
        "profiler: k-means execution time; Avg distance calculations ratio; total iterations; Best MSE (SSE). "
        "Logs: experiments/gap1_medium_runs.txt, experiments/gap3_runs.txt. "
        "Tables and plots generated with experiments/plot_pdc_figures.py."
    )

    doc.add_heading("8. Results", level=1)

    doc.add_heading("8.1 Gap 1 — Static-oriented vs guided (Elkan)", level=2)
    doc.add_paragraph(
        "Table 1 summarizes medians from paired STATIC/GUIDED blocks (see parse_gap1_log.py). "
        "Ratio guided/static uses comparable pairs only (same iterations and |ΔSSE|<0.5 on bracket total)."
    )
    t1 = doc.add_table(rows=5, cols=5)
    t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "Threads"
    t1.rows[0].cells[1].text = "Median static (s)"
    t1.rows[0].cells[2].text = "Median guided (s)"
    t1.rows[0].cells[3].text = "Median ratio G/S"
    t1.rows[0].cells[4].text = "Comparable / total"
    rows_a = [
        ("1", "18.9018", "18.8630", "0.9943", "5 / 5"),
        ("2", "11.3643", "10.9344", "0.9694", "5 / 5"),
        ("4", "7.9637", "7.9121", "—", "0 / 5"),
        ("8", "8.6657", "8.0045", "—", "0 / 5"),
    ]
    for i, row in enumerate(rows_a, start=1):
        for j, val in enumerate(row):
            t1.rows[i].cells[j].text = val
    doc.add_paragraph(
        "At T=4 and T=8, no pairs met the comparability criterion (DIFF_PATH). "
        "Do not claim a scheduling ‘winner’ from median times alone when paths differ."
    )

    for name, cap in (
        ("fig_gap1_elkan_static_vs_guided_median.png", "Figure 1. Gap 1 — median wall time (all pairs)."),
        ("fig_gap1_comparable_pairs_only.png", "Figure 2. Gap 1 — medians for comparable pairs only (T=1,2)."),
    ):
        fp = fig_dir / name
        if fp.is_file():
            doc.add_paragraph(cap)
            doc.add_picture(str(fp), width=Inches(6.2))
        else:
            doc.add_paragraph(f"[{cap} — file missing: {fp}]")

    doc.add_heading("8.2 Gap 3 — Dimensionality (N=100000, K=256, T=8)", level=2)
    doc.add_paragraph("Table 2. Runtime, pruning ratio, and SSE comparability notes.")
    t2 = doc.add_table(rows=7, cols=9)
    t2.style = "Table Grid"
    hdr2 = [
        "M",
        "Naive (s)",
        "Elk st",
        "Elk gd",
        "Ratio st",
        "Ratio gd",
        "it L/Es/Eg",
        "n vs Es",
        "Es vs Eg",
    ]
    for j, h in enumerate(hdr2):
        t2.rows[0].cells[j].text = h
    tb_rows = [
        ("2", "13.3161", "2.7458", "2.6948", "0.298", "0.298", "82/82/82", "match", "match"),
        ("8", "4.1285", "2.8219", "2.8476", "2.321", "2.432", "48/48/45", "match", "differs"),
        ("16", "4.7254", "3.2900", "2.8113", "5.671", "5.671", "47/47/47", "match", "match"),
        ("32", "4.1874", "2.5037", "2.4925", "11.716", "11.716", "30/30/30", "differs", "match"),
        ("64", "5.7816", "3.0624", "3.5013", "20.039", "19.459", "24/23/24", "differs", "differs"),
        ("128", "9.0140", "4.3204", "4.0998", "31.831", "31.791", "17/17/17", "differs", "differs"),
    ]
    for i, row in enumerate(tb_rows, start=1):
        for j, val in enumerate(row):
            t2.rows[i].cells[j].text = val

    for name, cap in (
        ("fig_gap3_ratio_vs_M.png", "Figure 3. Gap 3 — Avg distance calculations ratio vs M."),
        ("fig_gap3_time_vs_M.png", "Figure 4. Gap 3 — Wall time vs M (naive vs Elkan builds)."),
        ("fig_gap3_speedup_naive_over_elkan_static.png", "Figure 5. Gap 3 — Naive time / Elkan static time vs M."),
    ):
        fp = fig_dir / name
        if fp.is_file():
            doc.add_paragraph(cap)
            doc.add_picture(str(fp), width=Inches(6.2))
        else:
            doc.add_paragraph(f"[{cap} — file missing: {fp}]")

    doc.add_heading("9. Discussion", level=1)
    doc.add_paragraph(
        "Gap 1: At T=1–2, guided can be slightly faster when paths match; at T=4–8, scheduling-associated "
        "differences in reduction order yield mismatched iteration counts or SSE, so scheduling-only conclusions "
        "are ambiguous without deterministic reductions."
    )
    doc.add_paragraph(
        "Gap 3: The average distance-calculation ratio rises with M on uniform random data, consistent with "
        "weaker pruning when distances concentrate (high-dimensional phenomenon — cite Beyer et al. 1999; "
        "Aggarwal et al. 2001 in your PDF references list)."
    )

    doc.add_heading("9.1 Optional polish — variance and future work", level=2)
    doc.add_paragraph(
        "Single timed runs per Gap 3 configuration: repeat each configuration 3–5× and report median/IQR for "
        "production-grade conclusions. Clustered synthetic mixtures could better match K-means use cases than "
        "pure i.i.d. uniforms. Deterministic parallel reductions (if enabled in a fork) would reduce DIFF_PATH "
        "frequency when comparing schedules."
    )

    doc.add_heading("10. Conclusion", level=1)
    doc.add_paragraph(
        "We characterized Elkan vs naive across dimensions and compared OpenMP scheduling builds on one "
        "workstation. Scheduling effects are thread-dependent and sometimes non-comparable at higher T; pruning "
        "weakens as M increases on our synthetic generator. Future work: repeated trials, richer datasets, "
        "optional MPI scaling in lab environments."
    )

    doc.add_heading("11. LLM usage disclosure", level=1)
    doc.add_paragraph(
        "LLMs (e.g. Cursor/ChatGPT) assisted with outlining, environment troubleshooting (PEP 668, dos2unix), "
        "and plotting/log-parsing workflow suggestions. All timings, SSE values, and figures were produced locally "
        "from our binaries and logs; both authors reviewed results against stdout."
    )

    doc.add_heading("12. References (minimum)", level=1)
    for ref in (
        "W. Kwedlo & M. Czochanski, “A hybrid MPI/OpenMP parallelization of K-means algorithms accelerated using "
        "the triangle inequality,” IEEE Access, 2019. DOI 10.1109/ACCESS.2019.2907885.",
        "C. Elkan, “Using the triangle inequality to accelerate k-means,” ICML, 2003.",
        "G. Hamerly & C. Drake, Accelerating Lloyd’s Algorithm for K-Means Clustering, Springer, 2015.",
        "OpenMP Application Program Interface v4.5 (scheduling).",
        "K. Beyer et al., “When is nearest neighbor meaningful?” ICDT, 1999.",
        "C. Aggarwal et al., “On the surprising behavior of distance metrics in high dimensional space,” ICDT, 2001.",
        "hybrid-triangle-kmeans repository, Bitbucket, commit 0265d45a2eb04fec01ed53fc8635b277082ee284.",
    ):
        doc.add_paragraph(ref, style="List Bullet")

    doc.add_heading("Appendix A — Reproduction commands", level=1)
    add_mono_block(
        doc,
        """# Gap 1 summary (stdout pairs in gap1_medium_runs.txt)
python3 ~/22i-2327_F_FinalProject/experiments/parse_gap1_log.py \\
  ~/22i-2327_F_FinalProject/experiments/gap1_medium_runs.txt \\
  | tee ~/22i-2327_F_FinalProject/experiments/gap1_parsed_summary.txt

# Figures + CSV + TABLE_*.md
python3 ~/22i-2327_F_FinalProject/experiments/plot_pdc_figures.py \\
  --exp-dir ~/22i-2327_F_FinalProject/experiments

# Gap 3 matrix (after data_gap3 bins exist)
export OMP_NUM_THREADS=8
~/22i-2327_F_FinalProject/experiments/run_gap3_matrix.sh""",
    )

    doc.add_heading("Appendix B — Sample log headers (Gap 3)", level=1)
    if gap3_log.is_file():
        txt = gap3_log.read_text(encoding="utf-8", errors="replace")
        # First ~3500 chars of informative lines
        excerpt = "\n".join(
            ln
            for ln in txt.splitlines()
            if ln.startswith("========") or "k-means execution time:" in ln or "Avg distance" in ln or "Loading dataset" in ln
        )[:6000]
        if len(excerpt) > 5500:
            excerpt = excerpt[:5500] + "\n… [truncated]"
        add_mono_block(doc, excerpt or "(could not extract)")
    else:
        doc.add_paragraph("(gap3_runs.txt not found next to this script.)")

    doc.add_heading("Appendix C — Optional perf stat (example command)", level=1)
    doc.add_paragraph(
        "If linux-tools-generic / perf is available, a single aggregate hardware counter snapshot can complement "
        "program metrics. Example (edit paths if needed):"
    )
    add_mono_block(
        doc,
        """perf stat -e cycles,instructions,cache-references,cache-misses \\
  ~/22i-2327_F_FinalProject/experiments/bin/kmeans_static_gap1 \\
  ~/22i-2327_F_FinalProject/experiments/data_gap3/gap3_N100000_M32.bin \\
  -a elkan -c 256 -v 1 -r 42 -R 1e-4""",
    )
    perf_out = try_shell_capture(
        [
            "perf",
            "stat",
            "-e",
            "cycles,instructions,cache-references,cache-misses",
            str(project_root / "experiments/bin/kmeans_static_gap1"),
            str(project_root / "experiments/data_gap3/gap3_N100000_M32.bin"),
            "-a",
            "elkan",
            "-c",
            "256",
            "-v",
            "1",
            "-r",
            "42",
            "-R",
            "1e-4",
        ],
        timeout=120,
    )
    if perf_out:
        doc.add_paragraph("Captured perf output on this machine (best-effort under WSL):")
        add_mono_block(doc, perf_out[-4000:])
    else:
        doc.add_paragraph(
            "(perf not run or unavailable — run the command above on WSL and paste output here manually.)"
        )

    doc.add_heading("Appendix D — Gap 1 log excerpt", level=1)
    if gap1_log.is_file():
        head = try_shell_capture(["head", "-n", "80", str(gap1_log)])
        if head:
            add_mono_block(doc, head)
        else:
            add_mono_block(doc, gap1_log.read_text(encoding="utf-8", errors="replace")[:4000])
    else:
        doc.add_paragraph("(gap1_medium_runs.txt not found.)")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--project-root",
        type=Path,
        default=Path.home() / "22i-2327_F_FinalProject",
    )
    ap.add_argument(
        "--out",
        type=Path,
        default=None,
        help="Output .docx path",
    )
    args = ap.parse_args()
    root = args.project_root.expanduser().resolve()
    out = args.out or (root / "PDC_Kmeans_Final_Report.docx")
    build_doc(root, out)


if __name__ == "__main__":
    main()
